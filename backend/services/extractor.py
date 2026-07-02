"""
File parsing service – PDF and DOCX text extraction.

Libraries (verified latest stable):
  - pypdf 6.14.2  → PdfReader (PdfFileReader was removed in older releases)
  - python-docx 1.2.0 → Document() factory function

SECURITY: this module reads bytes and returns plain text.
It does NOT execute, evaluate, or interpret any content from the file.
"""

from __future__ import annotations

import io
import logging

from fastapi import HTTPException, status

logger = logging.getLogger(__name__)

# ── Allowed MIME types ────────────────────────────────────────────────────────
ALLOWED_CONTENT_TYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}

ALLOWED_EXTENSIONS = {".pdf", ".docx"}

# 5 MB hard limit
MAX_FILE_BYTES = 5 * 1024 * 1024


def validate_upload(filename: str, content_type: str, size: int) -> None:
    """
    Validate filename extension, MIME type, and size before parsing.
    Raises HTTP 400 on any violation.
    """
    ext = "." + filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Unsupported file type '{ext}'. Only PDF and DOCX are accepted.",
        )

    if content_type not in ALLOWED_CONTENT_TYPES:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=(
                f"Unacceptable Content-Type '{content_type}'. "
                "Upload a PDF (application/pdf) or DOCX file."
            ),
        )

    if size > MAX_FILE_BYTES:
        mb = size / (1024 * 1024)
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"File too large ({mb:.1f} MB). Maximum allowed size is 5 MB.",
        )


def extract_text_from_pdf(data: bytes) -> str:
    """
    Extract plain text from a PDF byte payload using pypdf 6.x PdfReader.

    - PdfFileReader is removed in pypdf 3+; we use PdfReader exclusively.
    - Returns concatenated page text joined by newlines.
    - Raises HTTP 422 if the file is encrypted or unreadable.
    """
    try:
        from pypdf import PdfReader  # pypdf 6.14.2
        from pypdf.errors import PdfReadError

        reader = PdfReader(io.BytesIO(data))

        if reader.is_encrypted:
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                detail="PDF is password-protected. Please upload an unlocked file.",
            )

        pages_text: list[str] = []
        for page in reader.pages:
            extracted = page.extract_text() or ""
            pages_text.append(extracted)

        raw = "\n".join(pages_text)

        if not raw.strip():
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                detail="Could not extract any text from the PDF. "
                       "It may be a scanned image PDF with no text layer.",
            )

        return raw

    except HTTPException:
        raise
    except Exception as exc:
        logger.warning("PDF extraction failed: %s", exc)
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="The PDF could not be read. It may be corrupted or malformed.",
        ) from exc


def extract_text_from_docx(data: bytes) -> str:
    """
    Extract plain text from a DOCX byte payload using python-docx 1.2.0.

    - Uses Document() factory function (stable since 0.3.0).
    - Iterates doc.paragraphs for body text.
    - Raises HTTP 422 if the file is unreadable.
    """
    try:
        from docx import Document  # python-docx 1.2.0

        doc = Document(io.BytesIO(data))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        raw = "\n".join(paragraphs)

        if not raw.strip():
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                detail="Could not extract any text from the DOCX. The document appears empty.",
            )

        return raw

    except HTTPException:
        raise
    except Exception as exc:
        logger.warning("DOCX extraction failed: %s", exc)
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="The DOCX could not be read. It may be corrupted or malformed.",
        ) from exc


def extract_text(filename: str, content_type: str, data: bytes) -> str:
    """
    Dispatch to the correct extractor based on file extension.
    Returns raw (unsanitized) text – caller must sanitize before use.
    """
    ext = "." + filename.rsplit(".", 1)[-1].lower()
    if ext == ".pdf":
        return extract_text_from_pdf(data)
    if ext == ".docx":
        return extract_text_from_docx(data)
    # Should not reach here after validate_upload(), but be defensive
    raise HTTPException(
        status_code=status.HTTP_400_BAD_REQUEST,
        detail=f"Unsupported extension '{ext}'.",
    )
